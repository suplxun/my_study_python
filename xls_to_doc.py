import openpyxl
from docx import Document
import os


os.chdir('C:\\Users\\Administrator\\Desktop')


wb = openpyxl.load_workbook('md.xlsx')
doc1 = Document()

sheet_a = wb["Sheet1"]

dict_names = {}
dict_head = {}
dict_preparation = {}
list_tea_res = []

for i in range(3,305):
        name = sheet_a.cell(row = i,column = 3).value
        grade = sheet_a.cell(row = i,column = 6).value
        course = sheet_a.cell(row = i,column = 7).value
        headteacher = sheet_a.cell(row = i,column = 9).value
        Leader_of_lesson_preparation = sheet_a.cell(row = i,column = 11).value
        Head_of_teaching_and_Research_Group = sheet_a.cell(row = i,column = 12).value
        if len(name) == 2:
                name = '  '.join(name)
        if (grade not in dict_names):
                dict_names[grade]={}
        if (course not in dict_names[grade]):
                dict_names[grade][course] = []
        dict_names[grade][course].append(name)

        if (grade not in dict_head):
                dict_head[grade]= []
        if(headteacher != None):
            dict_head[grade].append(name)

        if (grade not in dict_preparation):
            dict_preparation[grade]= []
        if(Leader_of_lesson_preparation != None):
            dict_preparation[grade].append(name)

        if (Head_of_teaching_and_Research_Group != None):
            list_tea_res.append(name)

s = ''
for g in dict_names.keys():
        s = s + g + '\n'
        for c in dict_names[g].keys():
                s = s + str(c) + '\n'
                count = len(dict_names[g][c])
                k = 0
                for n in dict_names[g][c]:
                    count = count - 1
                    if (k != 6):
                        s = s + n + "  "
                        k = k + 1
                        if (count == 0):
                            s = s + '\n'
                    elif (k == 6):
                        k = 0
                        s = s + n + "  " + '\n'
p=""
for g in dict_preparation.keys():
    p = p + g +"备课组长"+ '\n'
    count = len(dict_preparation[g])
    k=0
    for n in dict_preparation[g]:
        count = count -1
        if(k!=6):
            p = p+n+"  "
            k = k+1
            if(count == 0):
                p=p+"\n"
        elif(k==6):
            k=0
            p=p+n+"  "+"\n"

h=""
for g in dict_head.keys():
    h = h + g +"班主任"+ '\n'
    count = len(dict_head[g])
    k=0
    for n in dict_head[g]:
        count = count -1
        if(k!=6):
            h = h+n+"  "
            k = k+1
            if(count == 0):
                h=h+"\n"
        elif(k==6):
            k=0
            h=h+n+"  "+"\n"

r="教研组长\n"
k=0
count = len(list_tea_res)
for n in list_tea_res:
    count = count - 1
    if (k != 6):
        r = r + n + "  "
        k = k + 1
        if (count == 0):
            r = r + "\n"
    elif (k == 6):
        k = 0
        r = r + n + "  " + "\n"

doc1.add_paragraph(s)
doc1.add_paragraph(p)
doc1.add_paragraph(h)
doc1.add_paragraph(r)
doc1.save('d.docx')
                
